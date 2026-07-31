"""标题解析模块

支持从纯文本、Word(.docx)、Excel(.xlsx) 中提取文章标题列表。

所有公开函数签名统一为:
    parse_xxx(source, **kwargs) -> List[str]
返回清洗后的标题字符串列表。
"""

import re
import io
from typing import List, Optional
from pathlib import Path

from src.config import (
    MIN_TITLE_LENGTH,
    MAX_TITLE_LENGTH,
    SKIP_PATTERNS,
    TITLE_COLUMN_KEYWORDS,
)


def _row_values(row) -> List[str]:
    """兼容不同 openpyxl 版本，将一行单元格转为字符串列表。"""
    values = []
    for cell in row:
        value = getattr(cell, "value", cell)
        values.append(str(value) if value is not None else "")
    return values


def _is_likely_title(text: str) -> bool:
    """判断一段文本是否'像'一个文章标题。

    基于启发式规则：
    - 长度在合理范围内
    - 不包含明显的非标题特征
    - 不是纯数字或纯标点
    """
    text = text.strip()
    if not text:
        return False

    # 长度检查
    if len(text) < MIN_TITLE_LENGTH:
        return False
    if len(text) > MAX_TITLE_LENGTH:
        return False

    # 纯数字/纯标点
    if re.match(r'^[\d\s\.\-—、,，;；:：（）()\[\]{}#\*]+$', text):
        return False

    # 匹配跳过模式
    text_lower = text.lower().replace(" ", "")
    for pattern in SKIP_PATTERNS:
        if pattern.replace(" ", "").lower() in text_lower:
            return False

    return True


def _clean_title(text: str) -> str:
    """清洗标题：去除编号前缀、多余空白、特殊字符等。"""
    text = text.strip()

    # 去除行首编号，如 "1.", "1、", "(1)", "[1]", "1)", "1）"
    text = re.sub(r'^[\s]*[\[\(（]?\d+[\]\)）]?[\.\s、,，\t]+', '', text)

    # 合并多个空白为单个空格
    text = re.sub(r'\s+', ' ', text)

    # 去除首尾的引号、破折号等
    text = text.strip('"\'""''--—→→- ')

    return text.strip()


def parse_text(raw_text: str) -> List[str]:
    """从纯文本中提取标题列表。

    处理逻辑：
    1. 按换行分割
    2. 跳过空行和明显非标题的行
    3. 清洗每行标题（去编号、多余空白等）
    4. 去重（保留顺序）
    """
    lines = raw_text.strip().split('\n')
    titles: List[str] = []
    seen: set = set()

    for line in lines:
        cleaned = _clean_title(line)
        if not _is_likely_title(cleaned):
            continue
        # 去重（忽略大小写和空格差异）
        normalized = cleaned.lower().replace(" ", "")
        if normalized not in seen:
            seen.add(normalized)
            titles.append(cleaned)

    return titles


def _find_title_column(headers: List[str]) -> Optional[int]:
    """在表头列表中查找标题列的索引。

    通过关键词匹配（题名/标题/Title 等），返回第一个匹配的列索引。
    未找到返回 None。
    """
    for i, header in enumerate(headers):
        header_lower = header.strip().lower().replace(" ", "").replace("_", "")
        for keyword in TITLE_COLUMN_KEYWORDS:
            kw_clean = keyword.lower().replace(" ", "").replace("_", "")
            if kw_clean in header_lower:
                return i
    return None


def parse_docx(file_bytes: bytes) -> List[str]:
    """从 Word (.docx) 文件中提取标题。

    处理优先级：
    1. 先扫描表格，找到标题列 → 提取表格中的标题
    2. 如果没有表格或没找到标题列，提取段落文本作为候选标题
    """
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    titles: List[str] = []
    seen: set = set()

    # 步骤1：尝试从表格提取
    if doc.tables:
        for table in doc.tables:
            if not table.rows:
                continue
            # 取第一行作为表头
            header_cells = [cell.text for cell in table.rows[0].cells]
            title_col = _find_title_column(header_cells)

            if title_col is not None:
                for row in table.rows[1:]:  # 跳过表头行
                    if title_col < len(row.cells):
                        text = row.cells[title_col].text.strip()
                        if text and _is_likely_title(text):
                            cleaned = _clean_title(text)
                            normalized = cleaned.lower().replace(" ", "")
                            if normalized not in seen:
                                seen.add(normalized)
                                titles.append(cleaned)
                # 如果找到了标题列，表格模式优先，直接返回
                if titles:
                    return titles

    # 步骤2：从段落提取（后备方案）
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # 检查是否包含多个标题（以换行符分隔）
        for line in text.split('\n'):
            if _is_likely_title(line):
                cleaned = _clean_title(line)
                normalized = cleaned.lower().replace(" ", "")
                if normalized not in seen:
                    seen.add(normalized)
                    titles.append(cleaned)

    return titles


def parse_excel(file_bytes: bytes) -> List[str]:
    """从 Excel (.xlsx/.xls) 文件中提取标题。

    自动查找标题列（匹配表头关键词），然后按行提取。
    如果找不到标题列，使用第一列。
    """
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(file_bytes), read_only=True)
    titles: List[str] = []
    seen: set = set()

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]

        # 读取前 50 行来确定表头位置
        rows_preview = []
        for i, row in enumerate(ws.iter_rows()):
            rows_preview.append(_row_values(row))
            if i >= 50:
                break

        if not rows_preview:
            continue

        # 尝试在前 5 行中找表头
        title_col: Optional[int] = None
        header_row_idx = 0

        for i in range(min(5, len(rows_preview))):
            col = _find_title_column(rows_preview[i])
            if col is not None:
                title_col = col
                header_row_idx = i
                break

        # 如果没找到标题列，使用第一列
        if title_col is None:
            title_col = 0
            header_row_idx = 0

        # 提取标题：预览行只用于定位表头，真正提取时遍历完整 sheet
        for i, row in enumerate(ws.iter_rows()):
            if i <= header_row_idx:
                continue
            if title_col < len(row):
                cell_value = getattr(row[title_col], "value", row[title_col])
                if cell_value is None:
                    continue
                text = str(cell_value).strip()
                if text and _is_likely_title(text):
                    cleaned = _clean_title(text)
                    normalized = cleaned.lower().replace(" ", "")
                    if normalized not in seen:
                        seen.add(normalized)
                        titles.append(cleaned)

    wb.close()
    return titles


def parse_uploaded_file(file_bytes: bytes, filename: str) -> List[str]:
    """根据文件扩展名自动选择解析器。

    Args:
        file_bytes: 文件二进制内容
        filename: 文件名（用于判断扩展名）

    Returns:
        提取的标题列表

    Raises:
        ValueError: 不支持的文件格式
    """
    ext = Path(filename).suffix.lower()

    if ext == ".docx":
        return parse_docx(file_bytes)
    elif ext in (".xlsx", ".xls"):
        return parse_excel(file_bytes)
    elif ext == ".txt":
        text = file_bytes.decode("utf-8", errors="replace")
        return parse_text(text)
    else:
        raise ValueError(f"不支持的文件格式: {ext}，支持的格式: .docx, .xlsx, .xls, .txt")
